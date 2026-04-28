import os
from collections import defaultdict
from datetime import datetime, date
from flask import Flask, render_template, redirect, url_for, request, flash, send_file, abort
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash
from sqlalchemy import func
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def set_chinese_font(run, font_name='宋体', font_size=None):
    """设置run的中文字体"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)
    if font_size:
        run.font.size = Pt(font_size)

from models import db, User, Department, Overtime
from forms import LoginForm, RegisterForm, UserEditForm, UserCreateForm, DepartmentForm, OvertimeForm, OvertimeFilterForm

app = Flask(__name__)
app.config['SECRET_KEY'] = 'overtime-system-secret-key-2026'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///overtime.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db.init_app(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def init_db():
    db.create_all()
    if not Department.query.first():
        depts = ['教务处', '政教处', '总务处', '办公室', '教研组', '年级组']
        for name in depts:
            db.session.add(Department(name=name))
        db.session.commit()


@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))


@app.route('/setup', methods=['GET', 'POST'])
def setup():
    if User.query.filter_by(role='admin').first():
        return redirect(url_for('login'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        if username and password:
            admin = User(username=username, role='admin')
            admin.set_password(password)
            db.session.add(admin)
            db.session.commit()
            flash('管理员账号已创建，请登录', 'success')
            return redirect(url_for('login'))

    return '''
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>初始化设置 - 加班填报系统</title>
        <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
    </head>
    <body class="bg-light">
        <div class="container py-5">
            <div class="row justify-content-center">
                <div class="col-md-4">
                    <div class="card">
                        <div class="card-body">
                            <h4 class="text-center mb-4">创建管理员账号</h4>
                            <form method="POST">
                                <div class="mb-3">
                                    <label class="form-label">管理员用户名</label>
                                    <input type="text" name="username" class="form-control" required>
                                </div>
                                <div class="mb-3">
                                    <label class="form-label">管理员密码</label>
                                    <input type="password" name="password" class="form-control" required>
                                </div>
                                <button type="submit" class="btn btn-primary w-100">创建</button>
                            </form>
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    '''


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))

    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.check_password(form.password.data):
            login_user(user)
            return redirect(url_for('dashboard'))
        flash('用户名或密码错误', 'danger')
    return render_template('login.html', form=form)


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))

    form = RegisterForm()
    form.department_id.choices = [(0, '请选择部门')] + [(d.id, d.name) for d in Department.query.all()]

    if form.validate_on_submit():
        if form.department_id.data == 0:
            flash('请选择部门', 'warning')
            return render_template('register.html', form=form)

        user = User(username=form.username.data, role='teacher', department_id=form.department_id.data)
        user.set_password(form.password.data)
        db.session.add(user)
        db.session.commit()
        flash('注册成功，请登录', 'success')
        return redirect(url_for('login'))
    return render_template('register.html', form=form)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))


@app.route('/dashboard')
@login_required
def dashboard():
    my_count = Overtime.query.filter_by(user_id=current_user.id).count()
    my_hours = db.session.query(func.sum(Overtime.hours)).filter_by(user_id=current_user.id).scalar() or 0

    stats = {}
    if current_user.is_admin():
        stats['total_users'] = User.query.count()
        stats['total_records'] = Overtime.query.count()
        stats['total_hours'] = db.session.query(func.sum(Overtime.hours)).scalar() or 0
        stats['recent_records'] = Overtime.query.order_by(Overtime.created_at.desc()).limit(5).all()
    elif current_user.is_manager():
        dept_users = User.query.filter_by(department_id=current_user.department_id).all()
        user_ids = [u.id for u in dept_users]
        stats['dept_users'] = len(dept_users)
        stats['dept_records'] = Overtime.query.filter(Overtime.user_id.in_(user_ids)).count()
        stats['dept_hours'] = db.session.query(func.sum(Overtime.hours)).filter(Overtime.user_id.in_(user_ids)).scalar() or 0
        stats['recent_records'] = Overtime.query.filter(Overtime.user_id.in_(user_ids)).order_by(Overtime.created_at.desc()).limit(5).all()
    else:
        stats['recent_records'] = Overtime.query.filter_by(user_id=current_user.id).order_by(Overtime.created_at.desc()).limit(5).all()

    return render_template('index.html', stats=stats, my_count=my_count, my_hours=my_hours)


@app.route('/overtime')
@login_required
def overtime_list():
    page = request.args.get('page', 1, type=int)
    per_page = 20

    query = Overtime.query.filter_by(user_id=current_user.id)
    if current_user.is_admin():
        query = Overtime.query
    elif current_user.is_manager():
        dept_users = User.query.filter_by(department_id=current_user.department_id).all()
        user_ids = [u.id for u in dept_users]
        query = Overtime.query.filter(Overtime.user_id.in_(user_ids))

    pagination = query.order_by(Overtime.date.desc()).paginate(page=page, per_page=per_page, error_out=False)
    records = pagination.items

    return render_template('overtime/list.html', records=records, pagination=pagination)


@app.route('/overtime/add', methods=['GET', 'POST'])
@login_required
def overtime_add():
    form = OvertimeForm()
    if form.validate_on_submit():
        record = Overtime(
            user_id=current_user.id,
            content=form.content.data,
            date=form.date.data,
            hours=form.hours.data,
            is_workday=form.is_workday.data,
            memo=form.memo.data
        )
        db.session.add(record)
        db.session.commit()
        flash('加班记录已保存', 'success')
        return redirect(url_for('overtime_list'))
    form.date.data = date.today()
    return render_template('overtime/add.html', form=form)


@app.route('/overtime/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def overtime_edit(id):
    record = Overtime.query.get_or_404(id)
    if not current_user.can_edit_overtime(record):
        abort(403)

    form = OvertimeForm(obj=record)
    if form.validate_on_submit():
        record.content = form.content.data
        record.date = form.date.data
        record.hours = form.hours.data
        record.is_workday = form.is_workday.data
        record.memo = form.memo.data
        db.session.commit()
        flash('加班记录已更新', 'success')
        return redirect(url_for('overtime_list'))
    return render_template('overtime/edit.html', form=form, record=record)


@app.route('/overtime/delete/<int:id>', methods=['POST'])
@login_required
def overtime_delete(id):
    record = Overtime.query.get_or_404(id)
    if not current_user.can_edit_overtime(record):
        abort(403)
    db.session.delete(record)
    db.session.commit()
    flash('加班记录已删除', 'success')
    return redirect(url_for('overtime_list'))


@app.route('/overtime/export')
@login_required
def overtime_export():
    records = Overtime.query.filter_by(user_id=current_user.id).order_by(Overtime.date.desc()).all()

    doc = Document()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('加班详细记录')
    set_chinese_font(run, '宋体', 18)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'（{datetime.now().year}年{datetime.now().month}月）')
    set_chinese_font(run, '宋体', 12)

    doc.add_paragraph()

    # 汇总信息
    total_days = len(set(r.date for r in records))
    total_hours = sum(r.hours for r in records)
    p = doc.add_paragraph()
    run = p.add_run(f'加班总天数：{total_days}天    加班总时长：{total_hours}小时')
    set_chinese_font(run, '宋体', 11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # 创建表格 - 每行一条记录
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # 表头
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(['姓名', '部门', '加班日期', '加班时长', '工作日', '加班内容']):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[i].paragraphs[0].runs:
            set_chinese_font(run, '宋体', 11)

    # 填充数据行
    for r in records:
        row = table.add_row()
        row.cells[0].text = r.user.username
        row.cells[1].text = r.user.department.name if r.user.department else ''
        row.cells[2].text = r.date.strftime('%Y-%m-%d')
        row.cells[3].text = f'{r.hours}小时'
        row.cells[4].text = '是' if r.is_workday else '否'
        row.cells[5].text = r.content
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                set_chinese_font(run, '宋体', 10)

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph('部门领导签字：')
    for run in p.runs:
        set_chinese_font(run, '宋体', 12)

    filename = f'{current_user.username}——加班记录.docx'
    temp_path = os.path.join('/tmp', filename)
    doc.save(temp_path)

    return send_file(temp_path, as_attachment=True, download_name=filename)


@app.route('/admin/users')
@login_required
def admin_users():
    if not current_user.is_admin():
        abort(403)

    page = request.args.get('page', 1, type=int)
    per_page = 20
    pagination = User.query.order_by(User.created_at.desc()).paginate(page=page, per_page=per_page, error_out=False)
    users = pagination.items

    return render_template('admin/users.html', users=users, pagination=pagination)


@app.route('/admin/user/add', methods=['GET', 'POST'])
@login_required
def admin_user_add():
    if not current_user.is_admin():
        abort(403)

    form = UserCreateForm()
    form.department_id.choices = [(0, '请选择部门')] + [(d.id, d.name) for d in Department.query.all()]

    if form.validate_on_submit():
        if form.department_id.data == 0:
            flash('请选择部门', 'warning')
            return render_template('admin/user_add.html', form=form)

        user = User(username=form.username.data, role=form.role.data, department_id=form.department_id.data)
        user.set_password('123456')
        db.session.add(user)
        db.session.commit()
        flash('用户已创建', 'success')
        return redirect(url_for('admin_users'))

    return render_template('admin/user_add.html', form=form)


@app.route('/admin/user/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def admin_user_edit(id):
    if not current_user.is_admin():
        abort(403)

    user = User.query.get_or_404(id)
    form = UserEditForm(obj=user)
    form.department_id.choices = [(0, '无')] + [(d.id, d.name) for d in Department.query.all()]

    if form.validate_on_submit():
        user.username = form.username.data
        user.role = form.role.data
        user.department_id = form.department_id.data if form.department_id.data != 0 else None
        db.session.commit()
        flash('用户已更新', 'success')
        return redirect(url_for('admin_users'))

    return render_template('admin/user_edit.html', form=form, user=user)


@app.route('/admin/user/delete/<int:id>', methods=['POST'])
@login_required
def admin_user_delete(id):
    if not current_user.is_admin():
        abort(403)

    user = User.query.get_or_404(id)
    if user.username == 'admin':
        flash('不能删除管理员', 'danger')
        return redirect(url_for('admin_users'))

    db.session.delete(user)
    db.session.commit()
    flash('用户已删除', 'success')
    return redirect(url_for('admin_users'))


@app.route('/admin/departments')
@login_required
def admin_departments():
    if not current_user.is_admin():
        abort(403)

    departments = Department.query.all()
    return render_template('admin/departments.html', departments=departments)


@app.route('/admin/department/add', methods=['GET', 'POST'])
@login_required
def admin_department_add():
    if not current_user.is_admin():
        abort(403)

    form = DepartmentForm()
    if form.validate_on_submit():
        dept = Department(name=form.name.data)
        db.session.add(dept)
        db.session.commit()
        flash('部门已创建', 'success')
        return redirect(url_for('admin_departments'))

    return render_template('admin/department_add.html', form=form)


@app.route('/admin/department/delete/<int:id>', methods=['POST'])
@login_required
def admin_department_delete(id):
    if not current_user.is_admin():
        abort(403)

    dept = Department.query.get_or_404(id)
    if User.query.filter_by(department_id=id).first():
        flash('该部门下有用户，无法删除', 'danger')
        return redirect(url_for('admin_departments'))

    db.session.delete(dept)
    db.session.commit()
    flash('部门已删除', 'success')
    return redirect(url_for('admin_departments'))


@app.route('/admin/stats')
@login_required
def admin_stats():
    if not current_user.is_admin() and not current_user.is_manager():
        abort(403)

    form = OvertimeFilterForm()

    if current_user.is_manager():
        form.department_id.choices = [(0, '全部')] + [(current_user.department_id, current_user.department.name)]
        form.department_id.data = current_user.department_id
    else:
        form.department_id.choices = [(0, '全部')] + [(d.id, d.name) for d in Department.query.all()]

    form.user_id.choices = [(0, '全部人员')] + [(u.id, u.username) for u in User.query.all()]

    dept_id = request.args.get('department_id', 0, type=int)
    user_id = request.args.get('user_id', 0, type=int)
    start_date = request.args.get('start_date', type=str)
    end_date = request.args.get('end_date', type=str)

    query = Overtime.query

    if current_user.is_manager():
        dept_users = User.query.filter_by(department_id=current_user.department_id).all()
        user_ids = [u.id for u in dept_users]
        query = query.filter(Overtime.user_id.in_(user_ids))
    elif dept_id > 0:
        dept_users = User.query.filter_by(department_id=dept_id).all()
        user_ids = [u.id for u in dept_users]
        query = query.filter(Overtime.user_id.in_(user_ids))

    if user_id > 0:
        query = query.filter(Overtime.user_id == user_id)

    if start_date:
        try:
            start = datetime.strptime(start_date, '%Y-%m-%d').date()
            query = query.filter(Overtime.date >= start)
        except:
            pass

    if end_date:
        try:
            end = datetime.strptime(end_date, '%Y-%m-%d').date()
            query = query.filter(Overtime.date <= end)
        except:
            pass

    records = query.order_by(Overtime.date.desc()).all()

    summary = db.session.query(
        Overtime.user_id,
        User.username,
        func.sum(Overtime.hours).label('total_hours'),
        func.count(Overtime.id).label('total_count')
    ).join(User).filter(Overtime.id.in_([r.id for r in records])).group_by(Overtime.user_id, User.username).all()

    total_hours = sum(r.total_hours for r in summary) if summary else 0

    return render_template('admin/stats.html', form=form, records=records, summary=summary, total_hours=total_hours)


@app.route('/admin/export')
@login_required
def admin_export():
    if not current_user.is_admin() and not current_user.is_manager():
        abort(403)

    dept_id = request.args.get('department_id', 0, type=int)
    user_id = request.args.get('user_id', 0, type=int)
    start_date = request.args.get('start_date', type=str)
    end_date = request.args.get('end_date', type=str)

    query = Overtime.query

    if current_user.is_manager():
        dept_users = User.query.filter_by(department_id=current_user.department_id).all()
        user_ids = [u.id for u in dept_users]
        query = query.filter(Overtime.user_id.in_(user_ids))
    elif dept_id > 0:
        dept_users = User.query.filter_by(department_id=dept_id).all()
        user_ids = [u.id for u in dept_users]
        query = query.filter(Overtime.user_id.in_(user_ids))

    if user_id > 0:
        query = query.filter(Overtime.user_id == user_id)

    if start_date:
        try:
            start = datetime.strptime(start_date, '%Y-%m-%d').date()
            query = query.filter(Overtime.date >= start)
        except:
            pass

    if end_date:
        try:
            end = datetime.strptime(end_date, '%Y-%m-%d').date()
            query = query.filter(Overtime.date <= end)
        except:
            pass

    records = query.order_by(Overtime.date.desc()).all()

    # 生成 Word 文档
    doc = Document()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('加班详细记录')
    set_chinese_font(run, '宋体', 18)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'（{datetime.now().year}年{datetime.now().month}月）')
    set_chinese_font(run, '宋体', 12)

    doc.add_paragraph()

    # 按用户分组统计
    if user_id > 0:
        target_user = User.query.get(user_id)
        filename_prefix = f'{target_user.username}——加班记录' if target_user else '加班记录'
    else:
        filename_prefix = '加班记录'

    # 汇总信息
    total_days = len(set(r.date for r in records))
    total_hours = sum(r.hours for r in records)
    p = doc.add_paragraph()
    run = p.add_run(f'加班总天数：{total_days}天    加班总时长：{total_hours}小时')
    set_chinese_font(run, '宋体', 11)

    doc.add_paragraph()

    # 创建表格 - 每行一条记录
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # 表头
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(['姓名', '部门', '加班日期', '加班时长', '工作日', '加班内容']):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[i].paragraphs[0].runs:
            set_chinese_font(run, '宋体', 11)

    # 按用户分组
    user_data = defaultdict(list)
    for r in records:
        user_data[r.user_id].append(r)

    for uid, user_recs in user_data.items():
        user_obj = user_recs[0].user
        for r in user_recs:
            row = table.add_row()
            row.cells[0].text = user_obj.username
            row.cells[1].text = user_obj.department.name if user_obj.department else ''
            row.cells[2].text = r.date.strftime('%Y-%m-%d')
            row.cells[3].text = f'{r.hours}小时'
            row.cells[4].text = '是' if r.is_workday else '否'
            row.cells[5].text = r.content
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    set_chinese_font(run, '宋体', 10)

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph('部门领导签字：')
    for run in p.runs:
        set_chinese_font(run, '宋体', 12)

    filename = f'{filename_prefix}.docx'
    temp_path = os.path.join('/tmp', filename)
    doc.save(temp_path)

    return send_file(temp_path, as_attachment=True, download_name=filename)


if __name__ == '__main__':
    with app.app_context():
        init_db()
    app.run(debug=True, host='0.0.0.0', port=9000)
